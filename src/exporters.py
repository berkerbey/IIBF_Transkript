import os
from docx import Document
from src.utils import format_time, format_srt_time
from src.logger import logger

def export_txt(segments, output_path: str):
    """
    Export segments to a plain text file.
    Format is optimized for CAQDAS (MAXQDA / ATLAS.ti) which use timestamped text.
    """
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            for seg in segments:
                f.write(f"[{format_time(seg['start'])}] {seg['text']}\n")
        logger.info(f"Successfully exported TXT to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to export TXT: {str(e)}")
        return False

def export_docx(segments, output_path: str, metadata: dict = None):
    """
    Export segments to a DOCX file.
    Format is optimized for CAQDAS compatibility.
    """
    try:
        doc = Document()
        doc.add_heading("Transkript", level=1)

        if metadata:
            for key, value in metadata.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph("") # empty line

        table = None
        for seg in segments:
            p = doc.add_paragraph()
            p.add_run(f"[{format_time(seg['start'])}] ").bold = True
            p.add_run(seg["text"])

        doc.save(output_path)
        logger.info(f"Successfully exported DOCX to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to export DOCX: {str(e)}")
        return False

def export_srt(segments, output_path: str):
    """
    Export segments to an SRT file.
    """
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            for idx, seg in enumerate(segments, 1):
                f.write(f"{idx}\n")
                f.write(f"{format_srt_time(seg['start'])} --> {format_srt_time(seg['end'])}\n")
                f.write(f"{seg['text']}\n\n")
        logger.info(f"Successfully exported SRT to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to export SRT: {str(e)}")
        return False
